from flask import Flask, abort, redirect, request, render_template, send_from_directory, url_for, send_file
from io import BytesIO
import pdfplumber
import docx
import re
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import json
import os
import sqlite3
from datetime import datetime
from pathlib import Path
import google.generativeai as genai
from dotenv import load_dotenv
import time
from google.api_core import exceptions
from werkzeug.utils import secure_filename

load_dotenv()

genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

model = genai.GenerativeModel("models/gemini-2.5-flash")

app = Flask(__name__)
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
DATABASE = os.path.join(BASE_DIR, "hiring_dashboard.db")
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# Load spaCy model (run: python -m spacy download en_core_web_sm)
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    nlp = None

# Comprehensive technical skill regex patterns
TECH_PATTERNS = [

    # =========================
    # Programming Languages
    # =========================
    r'\bc\b', r'\bc\+\+\b', r'\bc#\b',
    r'\bpython\b', r'\bjava\b',
    r'\bjavascript\b', r'\btypescript\b',
    r'\bruby\b', r'\bgo\b', r'\brust\b',
    r'\bswift\b', r'\bkotlin\b',
    r'\bphp\b', r'\bperl\b',
    r'\br\b', r'\bmatlab\b',
    r'\bscala\b', r'\bdart\b',
    r'\bsql\b', r'\bnosql\b',

    # =========================
    # Web Development
    # =========================
    r'\bhtml\b', r'\bhtml5\b',
    r'\bcss\b', r'\bcss3\b',
    r'\bsass\b', r'\bscss\b',
    r'\bbootstrap\b', r'\btailwind\b',
    r'\bmaterial ui\b', r'\bchakra ui\b',

    r'\breact\b', r'\breact\.?js\b',
    r'\bnext\.?js\b',
    r'\bangular\b',
    r'\bvue\b', r'\bvue\.?js\b',
    r'\bnuxt\.?js\b',
    r'\bsvelte\b',
    r'\bjquery\b',
    r'\bredux\b',

    r'\bnode\.?js\b',
    r'\bexpress\.?js\b',
    r'\bnest\.?js\b',
    r'\bfastapi\b',
    r'\bflask\b',
    r'\bdjango\b',
    r'\blaravel\b',
    r'\bspring boot\b',

    r'\brest api\b',
    r'\bgraphql\b',
    r'\bjson\b',
    r'\bxml\b',

    # =========================
    # Mobile Development
    # =========================
    r'\bflutter\b',
    r'\breact native\b',
    r'\bandroid studio\b',
    r'\bxcode\b',
    r'\bionic\b',
    r'\bxamarin\b',

    # =========================
    # Databases
    # =========================
    r'\bmysql\b',
    r'\bpostgresql\b',
    r'\bmongodb\b',
    r'\bsqlite\b',
    r'\boracle\b',
    r'\bredis\b',
    r'\bcassandra\b',
    r'\bmariadb\b',
    r'\belasticsearch\b',
    r'\bfirebase\b',
    r'\bsupabase\b',

    # =========================
    # AI / Machine Learning
    # =========================
    r'\bmachine learning\b',
    r'\bdeep learning\b',
    r'\bdata science\b',
    r'\bnatural language processing\b',
    r'\bnlp\b',
    r'\bcomputer vision\b',
    r'\bneural network\b',
    r'\bcnn\b',
    r'\brnn\b',
    r'\blstm\b',
    r'\btransformer\b',
    r'\byolo\b',

    r'\btensorflow\b',
    r'\bpytorch\b',
    r'\bkeras\b',
    r'\bscikit[- ]learn\b',
    r'\bsklearn\b',
    r'\bpandas\b',
    r'\bnumpy\b',
    r'\bopencv\b',
    r'\bhugging face\b',
    r'\blangchain\b',
    r'\bopenai\b',
    r'\bchatgpt\b',
    r'\bllm\b',

    # =========================
    # Cloud / DevOps
    # =========================
    r'\baws\b',
    r'\bazure\b',
    r'\bgcp\b',
    r'\bdocker\b',
    r'\bkubernetes\b',
    r'\bterraform\b',
    r'\bansible\b',
    r'\bjenkins\b',
    r'\bgithub actions\b',
    r'\bgitlab ci\b',
    r'\bcircleci\b',
    r'\btravis ci\b',
    r'\bvagrant\b',
    r'\bhelm\b',

    r'\bgit\b',
    r'\bgithub\b',
    r'\bbitbucket\b',

    # =========================
    # Operating Systems
    # =========================
    r'\blinux\b',
    r'\bubuntu\b',
    r'\bunix\b',
    r'\bwindows server\b',
    r'\bmacos\b',

    # =========================
    # Cybersecurity
    # =========================
    r'\bkali linux\b',
    r'\bwireshark\b',
    r'\bnmap\b',
    r'\bmetasploit\b',
    r'\bburp suite\b',
    r'\bnessus\b',
    r'\bsplunk\b',
    r'\bsiem\b',

    r'\bethical hacking\b',
    r'\bpenetration testing\b',
    r'\bnetwork security\b',
    r'\bcybersecurity\b',
    r'\bfirewall\b',
    r'\bincident response\b',
    r'\bdigital forensics\b',

    # =========================
    # Networking
    # =========================
    r'\bnetworking\b',
    r'\btcp/ip\b',
    r'\bosi model\b',
    r'\brouter\b',
    r'\bswitch\b',
    r'\bcisco\b',
    r'\bpacket tracer\b',

    # =========================
    # Data Engineering / Big Data
    # =========================
    r'\bhadoop\b',
    r'\bspark\b',
    r'\bkafka\b',
    r'\bairflow\b',
    r'\betl\b',
    r'\bdata warehouse\b',
    r'\bsnowflake\b',
    r'\bdatabricks\b',

    # =========================
    # Data Analytics / BI
    # =========================
    r'\btableau\b',
    r'\bpower\s?bi\b',
    r'\bexcel\b',
    r'\bgoogle sheets\b',

    # =========================
    # Testing / QA
    # =========================
    r'\bselenium\b',
    r'\bcypress\b',
    r'\bjest\b',
    r'\bmocha\b',
    r'\bpytest\b',
    r'\bunit testing\b',
    r'\bautomation testing\b',

    # =========================
    # UI/UX & Creative Design
    # =========================
    r'\bfigma\b',
    r'\bsketch\b',
    r'\badobe xd\b',
    r'\bwireframing\b',
    r'\bprototyping\b',

    r'\badobe photoshop\b',
    r'\badobe illustrator\b',
    r'\badobe indesign\b',
    r'\badobe premiere\b',
    r'\bafter effects\b',
    r'\bcanva\b',
    r'\bcoreldraw\b',
    r'\bblender\b',
    r'\bmaya\b',

    # =========================
    # Engineering / CAD
    # =========================
    r'\bautocad\b',
    r'\bsolidworks\b',
    r'\bcatia\b',
    r'\bansys\b',
    r'\bsimulink\b',
    r'\bcreo\b',
    r'\binventor\b',
    r'\brevit\b',
    r'\barchicad\b',

    # =========================
    # Business / Finance / ERP
    # =========================
    r'\bsap\b',
    r'\boracle financials\b',
    r'\bquickbooks\b',
    r'\bzoho books\b',
    r'\bsalesforce\b',
    r'\bhubspot\b',
    r'\bcrm\b',
    r'\berp\b',
    r'\bservicenow\b',

    # =========================
    # Statistics / Research
    # =========================
    r'\bstata\b',
    r'\bspss\b',
    r'\bsas\b',
    r'\brapidminer\b',

    # =========================
    # Healthcare IT
    # =========================
    r'\bepic systems\b',
    r'\bmeditech\b',
    r'\belectronic medical records\b',
    r'\bemr\b',

    # =========================
    # Project Management
    # =========================
    r'\bjira\b',
    r'\bconfluence\b',
    r'\btrello\b',
    r'\basana\b',
    r'\bslack\b',
    r'\bnotion\b',
    r'\bmonday\.com\b',
    r'\bclickup\b',
    r'\bmiro\b',
    r'\bzoom\b',
    r'\bmicrosoft teams\b',

    # =========================
    # Game Development
    # =========================
    r'\bunity\b',
    r'\bunreal engine\b',
    r'\bgodot\b',

    # =========================
    # IoT / Embedded Systems
    # =========================
    r'\barduino\b',
    r'\braspberry pi\b',
    r'\besp32\b',
    r'\bembedded systems\b',

    # =========================
    # General Technical Concepts
    # =========================
    r'\boop\b',
    r'\bobject oriented programming\b',
    r'\bdata structures\b',
    r'\balgorithms\b',
    r'\bsoftware engineering\b',
    r'\bsdlc\b',
    r'\bagile\b',
    r'\bscrum\b',
    r'\bapi development\b',
    r'\bmicroservices\b'
]


def get_db_connection():
    conn = sqlite3.connect(DATABASE)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    with get_db_connection() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS jobs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                company TEXT NOT NULL,
                description TEXT NOT NULL,
                created_at TEXT NOT NULL
            )
            """
        )
        columns = [row["name"] for row in conn.execute("PRAGMA table_info(jobs)").fetchall()]
        if "active" not in columns:
            conn.execute("ALTER TABLE jobs ADD COLUMN active INTEGER NOT NULL DEFAULT 1")
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS candidate_scores (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                job_id INTEGER NOT NULL,
                candidate_name TEXT NOT NULL,
                file_name TEXT NOT NULL,
                score REAL NOT NULL,
                matched_skills TEXT NOT NULL,
                missing_skills TEXT NOT NULL,
                created_at TEXT NOT NULL,
                FOREIGN KEY (job_id) REFERENCES jobs (id)
            )
            """
        )
        score_columns = [row["name"] for row in conn.execute("PRAGMA table_info(candidate_scores)").fetchall()]
        if "summary" not in score_columns:
            conn.execute("ALTER TABLE candidate_scores ADD COLUMN summary TEXT NOT NULL DEFAULT ''")
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS activity (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                job_id INTEGER,
                message TEXT NOT NULL,
                score REAL,
                created_at TEXT NOT NULL,
                FOREIGN KEY (job_id) REFERENCES jobs (id)
            )
            """
        )

init_db()

def candidate_name_from_filename(filename):
    stem = Path(filename).stem
    cleaned = re.sub(r"[_-]+", " ", stem)
    cleaned = re.sub(r"\b(resume|cv|profile)\b", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned.title() if cleaned else filename

def save_uploaded_resume(file):
    original_filename = secure_filename(file.filename) or "resume"
    stem = Path(original_filename).stem or "resume"
    suffix = Path(original_filename).suffix.lower()
    timestamp = datetime.utcnow().strftime("%Y%m%d%H%M%S%f")
    saved_filename = f"{stem}_{timestamp}{suffix}"
    file.stream.seek(0)
    file.save(os.path.join(UPLOAD_DIR, saved_filename))
    file.stream.seek(0)
    return original_filename, saved_filename

def match_strength(score):
    if score >= 85:
        return "Strong match"
    if score >= 70:
        return "Good match"
    if score >= 50:
        return "Moderate match"
    return "Needs review"

def generate_ai_summary(candidate_name, job_title, job_description, matched, missing):
    
    prompt = f"""
    Role: {job_title}
    Target Requirements: {job_description}
    
    Candidate: {candidate_name}
    Matched Skills/Experience: {", ".join(matched)}
    Missing/Unclear Areas: {", ".join(missing)}

    Task: Write a 3-4 sentence professional recruiter "assessment" summary. 
    
    Guidelines:
    1. Start with a clear verdict (e.g., "[Name] is a very strong fit" or "A solid technical match").
    2. Synthesize the matched skills into career themes (e.g., instead of saying "knows Python," say "strong technical foundation in backend automation").
    3. Call out specific high-value areas (like Design Systems, User Research, or Prototyping).
    4. Mention "minor uncertainties" or gaps in a nuanced way.
    5. Avoid bullet points; use a fluid, professional narrative tone exactly like a recruiter speaking to a hiring manager.
    """

    # Inside your loop
    try:
            response = model.generate_content(prompt)
            return response.text.strip()
            
    except exceptions.ResourceExhausted:
            # If quota is hit, return this string instead of crashing
            print("Quota exceeded, skipping AI summary for now.")
            return "Summary pending (Rate limit reached). Please refresh later."
            
    except Exception as e:
            # Catch other errors (like internet issues)
            print(f"An error occurred: {e}")
            return "Summary currently unavailable."

#Job seeker: helper function
def check_ats_structure(resume_text):
    text = resume_text.lower()

    standard_sections = {
        "Skills": ["skills", "technical skills"],
        "Education": ["education", "academic background"],
        "Experience": ["experience", "work experience", "employment history"],
        "Projects": ["projects", "project experience"],
        "Certifications": ["certifications", "certificates"]
    }

    feedback = []

    for section, keywords in standard_sections.items():
        found = any(keyword in text for keyword in keywords)

        if found:
            feedback.append({
                "section": section,
                "status": "Detected",
                "message": f"{section} section is detected."
            })
        else:
            feedback.append({
                "section": section,
                "status": "Missing",
                "message": f"{section} section is not clearly detected. Consider adding this section."
            })

    return feedback

#Job seeker: AI Feedback Function
def generate_job_seeker_ai_feedback(score, matched, missing, ats_feedback, job_description):
    missing_sections = [
        item["section"] for item in ats_feedback if item["status"] == "Missing"
    ]

    prompt = f"""
    You are an AI resume advisor helping a job seeker improve their resume for a target job.

    Target Job Description:
    {job_description}

    Resume Match Score: {score}%

    Matched Skills:
    {", ".join(matched) if matched else "No strong matched skills detected."}

    Missing Skills:
    {", ".join(missing) if missing else "No major missing skills detected."}

    Missing ATS Sections:
    {", ".join(missing_sections) if missing_sections else "No major missing ATS sections detected."}

    Task:
    Write a clear and honest resume analysis for the job seeker.

    Requirements:
    1. Start by explaining what the match score means.
    2. Mention the candidate's strongest matching areas.
    3. Clearly explain the missing skills or skill gaps.
    4. Mention ATS structure issues if any.
    5. Give practical improvement suggestions.
    6. Do not exaggerate. Be honest but encouraging.
    7. Write in 2 short paragraphs, suitable for a student/job seeker.
    """

    try:
        response = model.generate_content(prompt)
        return response.text.strip()

    except exceptions.ResourceExhausted:
        return "AI feedback is temporarily unavailable due to usage limits. However, you can improve your resume by adding missing skills, using clear ATS-friendly section headings, and aligning your wording with the job description."

    except Exception as e:
        print(f"AI feedback error: {e}")
        return "AI feedback is currently unavailable. Please review the matched skills, missing skills, and ATS section feedback below."

def format_activity_date(value):
    try:
        return datetime.fromisoformat(value).strftime("%d %b")
    except ValueError:
        return value

def build_recruiter_overview():
    with get_db_connection() as conn:
        jobs = conn.execute(
            """
            SELECT
                j.id,
                j.title,
                j.company,
                j.description,
                j.created_at,
                COUNT(cs.id) AS total_candidates,
                MAX(cs.score) AS top_score,
                AVG(cs.score) AS average_score
            FROM jobs j
            LEFT JOIN candidate_scores cs ON cs.job_id = j.id
            WHERE j.active = 1
            GROUP BY j.id
            ORDER BY j.id DESC
            """
        ).fetchall()
        total_candidates = conn.execute("SELECT COUNT(*) AS total FROM candidate_scores").fetchone()["total"]
        average_match = conn.execute("SELECT AVG(score) AS average FROM candidate_scores").fetchone()["average"]
        top_candidate = conn.execute(
            """
            SELECT cs.candidate_name, cs.score, j.title
            FROM candidate_scores cs
            JOIN jobs j ON j.id = cs.job_id
            ORDER BY cs.score DESC, cs.id DESC
            LIMIT 1
            """
        ).fetchone()
        activities = conn.execute(
            """
            SELECT message, score, created_at
            FROM activity
            ORDER BY id DESC
            LIMIT 8
            """
        ).fetchall()

    job_cards = []
    for job in jobs:
        skills = sorted(extract_required_skills(job["description"]))[:4]
        job_cards.append({
            "id": job["id"],
            "title": job["title"],
            "company": job["company"],
            "skills": skills,
            "total_candidates": job["total_candidates"],
            "top_score": round(job["top_score"]) if job["top_score"] is not None else 0,
            "average_score": round(job["average_score"]) if job["average_score"] is not None else 0,
        })

    if top_candidate:
        top_candidate_card = {
            "name": top_candidate["candidate_name"],
            "score": round(top_candidate["score"]),
            "role": top_candidate["title"],
            "is_demo": False,
        }
    else:
        top_candidate_card = {
            "name": "Aiyana Brooks",
            "score": 92,
            "role": "Senior Full-Stack Engineer",
            "is_demo": True,
        }

    return {
        "open_postings": len(job_cards),
        "candidates_scored": total_candidates,
        "average_match": round(average_match) if average_match is not None else 0,
        "top_candidate": top_candidate_card,
        "activities": [
            {
                "message": activity["message"],
                "score": round(activity["score"]) if activity["score"] is not None else None,
                "date": format_activity_date(activity["created_at"]),
            }
            for activity in activities
        ],
    }, job_cards

def get_job_metrics(job_id):
    with get_db_connection() as conn:
        metrics = conn.execute(
            """
            SELECT
                COUNT(id) AS total_candidates,
                MAX(score) AS top_score,
                AVG(score) AS average_score
            FROM candidate_scores
            WHERE job_id = ?
            """,
            (job_id,),
        ).fetchone()

    return {
        "total_candidates": metrics["total_candidates"],
        "top_score": round(metrics["top_score"]) if metrics["top_score"] is not None else 0,
        "average_score": round(metrics["average_score"]) if metrics["average_score"] is not None else 0,
    }

def get_ranked_candidates(job_id):
    with get_db_connection() as conn:
        # 1. Fetch Job Details first so the AI has context
        job = conn.execute(
            "SELECT title, description FROM jobs WHERE id = ?", 
            (job_id,)
        ).fetchone()
        
        # Fallback if job is missing
        job_title = job["title"] if job else "Unknown Position"
        job_desc = job["description"] if job else ""

        # 2. Fetch Candidate Scores
        rows = conn.execute(
            """
            SELECT id, candidate_name, file_name, score, matched_skills, missing_skills, summary
            FROM candidate_scores
            WHERE job_id = ?
            ORDER BY score DESC, id DESC
            """,
            (job_id,),
        ).fetchall()

    candidates = []
    for row in rows:
        matched = json.loads(row["matched_skills"] or "[]")
        missing = json.loads(row["missing_skills"] or "[]")
        
        # 3. Update the function call to include all 5 arguments
        summary = row["summary"]
        if not summary:
            summary = generate_ai_summary(
                candidate_name=row["candidate_name"],
                job_title=job_title,
                job_description=job_desc,
                matched=matched,
                missing=missing
            )

        candidates.append({
            "id": row["id"],
            "candidate_name": row["candidate_name"],
            "file_name": row["file_name"],
            "score": round(row["score"]),
            "matched": matched,
            "missing": missing,
            "summary": summary,
            "strength": match_strength(row["score"]),
        })
    return candidates

def extract_required_skills(job_description):
    text = job_description.lower()
    found_skills = set()

    # Regex-based detection
    for pattern in TECH_PATTERNS:
        matches = re.findall(pattern, text)
        found_skills.update(matches)

    # Optional: spaCy NER for extra tech/product mentions
    if nlp:
        doc = nlp(job_description)
        for ent in doc.ents:
            if ent.label_ in ("ORG", "PRODUCT","SKILL"):
                found_skills.add(ent.text.lower())

    return list(found_skills)

def parse_pdf(file):
    text = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                text += page.extract_text() + " "
    return text.lower()

def parse_docx(file):
    doc = docx.Document(file)
    text = " ".join([para.text for para in doc.paragraphs])
    return text.lower()

def calculate_combined_score(resume_text, job_description, required_skills):
    # Keyword match score
    matched = sorted(skill for skill in required_skills if skill in resume_text)
    missing = sorted(skill for skill in required_skills if skill not in resume_text)
    keyword_score = (len(matched) / len(required_skills)) * 100 if required_skills else 0

    # TF-IDF similarity score
    documents = [resume_text, job_description.lower()]
    vectorizer = TfidfVectorizer(stop_words="english")
    tfidf_matrix = vectorizer.fit_transform(documents)
    tfidf_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0] * 100

    # Combine both into one score (tune weights as needed)
    combined_score = (0.6 * keyword_score) + (0.4 * tfidf_score)

    return round(combined_score, 2), matched, missing

def safe_json_loads(value):
    try:
        return json.loads(value) if value else []
    except json.JSONDecodeError:
        return []


def format_list(items, key):
    values = []
    for item in items:
        text = item.get(key, "").strip()
        if text:
            values.append(text)
    return values


def build_resume_text_from_builder(resume_data):
    resume_parts = []

    resume_parts.append(resume_data.get("summary", ""))

    for skill in resume_data.get("skills", []):
        resume_parts.append(skill)

    for exp in resume_data.get("experiences", []):
        resume_parts.append(exp.get("role", ""))
        resume_parts.append(exp.get("company", ""))
        resume_parts.append(exp.get("description", ""))

    for edu in resume_data.get("educations", []):
        resume_parts.append(edu.get("school", ""))
        resume_parts.append(edu.get("degree", ""))

    for project in resume_data.get("projects", []):
        resume_parts.append(project.get("title", ""))
        resume_parts.append(project.get("description", ""))
        resume_parts.append(project.get("tools", ""))

    for cert in resume_data.get("certifications", []):
        resume_parts.append(cert.get("name", ""))
        resume_parts.append(cert.get("provider", ""))

    return " ".join(resume_parts).lower()


def generate_ai_bullet_points(experience):
    prompt = f"""
    You are an AI resume writing assistant.

    Convert the following work experience into 3 professional resume bullet points.

    Company: {experience.get("company")}
    Role: {experience.get("role")}
    Period: {experience.get("period")}
    What the candidate did:
    {experience.get("description")}

    Requirements:
    1. Use strong action verbs.
    2. Make the bullets suitable for an ATS-friendly resume.
    3. Do not invent fake achievements or numbers.
    4. Keep each bullet concise.
    """

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print(f"AI bullet generation error: {e}")
        return experience.get("description", "")


def generate_ai_resume_summary(resume_data):
    prompt = f"""
    You are an AI resume writing assistant.

    Candidate Information:
    Name: {resume_data.get("full_name")}
    Target Role: {resume_data.get("target_role")}
    Current Summary: {resume_data.get("summary")}
    Skills: {", ".join(resume_data.get("skills", []))}
    Experience: {resume_data.get("experience_text")}
    Education: {resume_data.get("education_text")}
    Projects: {resume_data.get("project_text")}
    Certifications: {resume_data.get("certification_text")}

    Task:
    Rewrite the professional summary into an ATS-friendly resume summary.

    Requirements:
    1. Keep it around 2-3 sentences.
    2. Use professional language.
    3. Highlight relevant skills, education, projects, or experience.
    4. Do not invent fake experience, fake skills, or fake certifications.
    5. Make it suitable for the target role.
    """

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print(f"AI resume summary error: {e}")
        return resume_data.get("summary", "")


def generate_builder_ai_feedback(score, matched, missing, ats_feedback, target_role):
    missing_sections = [
        item["section"] for item in ats_feedback if item["status"] == "Missing"
    ]

    prompt = f"""
    You are an AI resume advisor.

    Target Role: {target_role}
    Resume Score: {score}%

    Matched Skills:
    {", ".join(matched) if matched else "No strong matched skills detected."}

    Missing Skills:
    {", ".join(missing) if missing else "No major missing skills detected."}

    Missing ATS Sections:
    {", ".join(missing_sections) if missing_sections else "No major missing sections detected."}

    Task:
    Give a clear explanation of the resume score and suggest what the job seeker should improve.

    Requirements:
    1. Explain what the score means.
    2. Mention strengths.
    3. Mention missing skills or skill gaps.
    4. Mention ATS structure issues if any.
    5. Give practical improvement suggestions.
    6. Use a helpful and honest tone.
    7. Write in 2 short paragraphs.
    """

    try:
        response = model.generate_content(prompt)
        return response.text.strip()
    except Exception as e:
        print(f"AI builder feedback error: {e}")
        return "AI feedback is currently unavailable. Please review the score, matched skills, missing skills, and ATS structure feedback."

@app.route("/", methods=["GET"])
def home():
    return render_template("index.html")

@app.route("/rank", methods=["POST"])
def rank_resumes():
    job_description = request.form.get("job_description", "")
    required_skills = extract_required_skills(job_description)
    submitted_names = [
        name.strip()
        for field_name in ("candidate_names", "candidate_name", "name", "full_name")
        for name in request.form.getlist(field_name)
        if name.strip()
    ]

    results = []
    for file in request.files.getlist("resumes"):
        submitted_filename = file.filename.lower()
        if not (submitted_filename.endswith(".pdf") or submitted_filename.endswith(".docx")):
            continue

        original_filename, saved_filename = save_uploaded_resume(file)
        if original_filename.lower().endswith(".pdf"):
            text = parse_pdf(file)
        elif original_filename.lower().endswith(".docx"):
            text = parse_docx(file)

        submitted_name = submitted_names.pop(0) if submitted_names else ""
        candidate_name = submitted_name or candidate_name_from_filename(original_filename)

        score, matched, missing = calculate_combined_score(text, job_description, required_skills)
        results.append({
            "candidate": candidate_name,
            "file_name": original_filename,
            "file_url": url_for("uploaded_file", filename=saved_filename),
            "score": score,
            "matched": matched,
            "missing": missing,
            "shortlisted": score >= 70,
        })

    results.sort(key=lambda x: x["score"], reverse=True)
    return render_template("result.html", results=results, required_skills=sorted(required_skills))


@app.route("/recruiter", methods=["GET", "POST"])
def recruiter():
    if request.method == "POST":
        title = request.form.get("title", "").strip()
        company = request.form.get("company", "").strip()
        description = request.form.get("description", "").strip()

        if title and company and description:
            with get_db_connection() as conn:
                cursor = conn.execute(
                    """
                    INSERT INTO jobs (title, company, description, created_at, active)
                    VALUES (?, ?, ?, ?, 1)
                    """,
                    (title, company, description, datetime.utcnow().isoformat(timespec="seconds")),
                )
                job_id = cursor.lastrowid
                conn.execute(
                    "INSERT INTO activity (job_id, message, score, created_at) VALUES (?, ?, ?, ?)",
                    (
                        job_id,
                        f"{title} opened at {company}",
                        None,
                        datetime.utcnow().isoformat(timespec="seconds"),
                    ),
                )
            return redirect(url_for("job_detail", job_id=job_id))

    kpis, jobs = build_recruiter_overview()

    return render_template("recruiter.html", jobs=jobs, kpis=kpis)

@app.route("/recruiter/jobs/<int:job_id>")
def job_detail(job_id):
    with get_db_connection() as conn:
        job = conn.execute("SELECT * FROM jobs WHERE id = ? AND active = 1", (job_id,)).fetchone()

    if job is None:
        abort(404)

    required_skills = sorted(extract_required_skills(job["description"]))
    return render_template(
        "recruiter.html",
        jobs=[],
        selected_job=job,
        required_skills=required_skills,
        job_metrics=get_job_metrics(job_id),
        ranked_candidates=get_ranked_candidates(job_id),
    )

@app.route("/recruiter/jobs/<int:job_id>/edit", methods=["POST"])
def edit_job(job_id):
    title = request.form.get("title", "").strip()
    company = request.form.get("company", "").strip()
    description = request.form.get("description", "").strip()

    if not title or not company or not description:
        return redirect(url_for("job_detail", job_id=job_id))

    with get_db_connection() as conn:
        job = conn.execute("SELECT * FROM jobs WHERE id = ? AND active = 1", (job_id,)).fetchone()
        if job is None:
            abort(404)

        timestamp = datetime.utcnow().isoformat(timespec="seconds")
        conn.execute(
            "UPDATE jobs SET title = ?, company = ?, description = ? WHERE id = ?",
            (title, company, description, job_id),
        )
        conn.execute(
            "INSERT INTO activity (job_id, message, score, created_at) VALUES (?, ?, ?, ?)",
            (job_id, f"{title} job description updated", None, timestamp),
        )

    return redirect(url_for("job_detail", job_id=job_id))

@app.route("/recruiter/jobs/<int:job_id>/delete", methods=["POST"])
def delete_job(job_id):
    with get_db_connection() as conn:
        job = conn.execute("SELECT * FROM jobs WHERE id = ? AND active = 1", (job_id,)).fetchone()
        if job is None:
            abort(404)

        timestamp = datetime.utcnow().isoformat(timespec="seconds")
        conn.execute("DELETE FROM candidate_scores WHERE job_id = ?", (job_id,))
        conn.execute("UPDATE jobs SET active = 0 WHERE id = ?", (job_id,))
        conn.execute(
            "INSERT INTO activity (job_id, message, score, created_at) VALUES (?, ?, ?, ?)",
            (job_id, f"{job['title']} archived and candidate data cleared", None, timestamp),
        )

    return redirect(url_for("recruiter"))

@app.route("/recruiter/jobs/<int:job_id>/candidate/<int:candidate_id>/delete", methods=["POST"])
def delete_candidate(job_id, candidate_id):
    with get_db_connection() as conn:
        candidate = conn.execute(
            """
            SELECT cs.candidate_name, cs.score, j.title
            FROM candidate_scores cs
            JOIN jobs j ON j.id = cs.job_id
            WHERE cs.job_id = ? AND cs.id = ? AND j.active = 1
            """,
            (job_id, candidate_id),
        ).fetchone()

        if candidate is None:
            abort(404)

        conn.execute("DELETE FROM candidate_scores WHERE job_id = ? AND id = ?", (job_id, candidate_id))

        timestamp = datetime.utcnow().isoformat(timespec="seconds")
        conn.execute(
            "INSERT INTO activity (job_id, message, score, created_at) VALUES (?, ?, ?, ?)",
            (
                job_id,
                f"{candidate['candidate_name']} removed from {candidate['title']}",
                candidate["score"],
                timestamp,
            ),
        )

    return redirect(url_for("job_detail", job_id=job_id))

@app.route("/recruiter/jobs/<int:job_id>/rank", methods=["POST"])
def rank_job_resumes(job_id):
    with get_db_connection() as conn:
        job = conn.execute("SELECT * FROM jobs WHERE id = ?", (job_id,)).fetchone()

    if job is None:
        abort(404)

    # 1. ALIGN VARIABLE NAMES
    # We extract these here so they are available for generate_ai_summary
    job_title = job["title"] 
    job_description = job["description"]
    
    required_skills = extract_required_skills(job_description)
    submitted_names = [name.strip() for name in request.form.getlist("candidate_names")]

    results = []
    timestamp = datetime.utcnow().isoformat(timespec="seconds")
    
    for index, file in enumerate(request.files.getlist("resumes")):
        filename = file.filename.lower()
        if not (filename.endswith(".pdf") or filename.endswith(".docx")):
            continue

        original_filename, saved_filename = save_uploaded_resume(file)
        filename = original_filename.lower()
        if filename.endswith(".pdf"):
            text = parse_pdf(file)
        elif filename.endswith(".docx"):
            text = parse_docx(file)

        score, matched, missing = calculate_combined_score(text, job_description, required_skills)
        
        # Determine candidate name
        candidate_name = submitted_names[index] if index < len(submitted_names) and submitted_names[index] else candidate_name_from_filename(original_filename)
        
        # 2. UPDATED FUNCTION CALL
        # Note: We use job_description (the variable we created above)
        summary = generate_ai_summary(
            candidate_name=candidate_name, 
            job_title=job_title, 
            job_description=job_description, 
            matched=matched, 
            missing=missing
        )

        results.append({
            "candidate": candidate_name,
            "candidate_name": candidate_name,
            "file_name": original_filename,
            "file_url": url_for("uploaded_file", filename=saved_filename),
            "score": score,
            "matched": matched,
            "missing": missing,
            "summary": summary,
            "shortlisted": score >= 70,
        })

    results.sort(key=lambda x: x["score"], reverse=True)
    
    if results:
        with get_db_connection() as conn:
            for result in results:
                conn.execute(
                    """
                    INSERT INTO candidate_scores 
                        (job_id, candidate_name, file_name, score, matched_skills, missing_skills, summary, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        job_id,
                        result["candidate_name"],
                        result["file_name"],
                        result["score"],
                        json.dumps(result["matched"]),
                        json.dumps(result["missing"]),
                        result["summary"],
                        timestamp,
                    ),
                )
                conn.execute(
                    "INSERT INTO activity (job_id, message, score, created_at) VALUES (?, ?, ?, ?)",
                    (
                        job_id,
                        f"{result['candidate_name']} added to {job_title}",
                        result["score"],
                        timestamp,
                    ),
                )
            conn.commit() # Don't forget to commit your changes!

    return render_template(
        "result.html",
        results=results,
        required_skills=sorted(required_skills),
        job=job,
    )

@app.route("/job_seeker")
def job_seeker():
    return render_template("job_seeker.html")

@app.route("/resume_builder")
def resume_builder():
    return render_template("resume_builder.html")

@app.route("/analyze_job_seeker_resume", methods=["POST"])
def analyze_job_seeker_resume():
    job_description = request.form.get("job_description", "").strip()
    resume_file = request.files.get("resume")

    if not resume_file or not job_description:
        return render_template(
            "job_seeker_result.html",
            error="Please upload a resume and paste a target job description."
        )

    filename = resume_file.filename.lower()

    if filename.endswith(".pdf"):
        resume_text = parse_pdf(resume_file)
    elif filename.endswith(".docx"):
        resume_text = parse_docx(resume_file)
    else:
        return render_template(
            "job_seeker_result.html",
            error="Unsupported file format. Please upload PDF or DOCX."
        )

    required_skills = extract_required_skills(job_description)

    score, matched, missing = calculate_combined_score(
        resume_text,
        job_description,
        required_skills
    )

    ats_feedback = check_ats_structure(resume_text)

    ai_feedback = generate_job_seeker_ai_feedback(
        score=score,
        matched=matched,
        missing=missing,
        ats_feedback=ats_feedback,
        job_description=job_description
    )

    return render_template(
        "job_seeker_result.html",
        score=score,
        matched=matched,
        missing=missing,
        required_skills=sorted(required_skills),
        ats_feedback=ats_feedback,
        ai_feedback=ai_feedback
    )

def check_builder_ats_structure(resume_data):
    feedback = []

    skills = resume_data.get("skills", [])
    educations = resume_data.get("educations", [])
    experiences = resume_data.get("experiences", [])
    projects = resume_data.get("projects", [])
    certifications = resume_data.get("certifications", [])

    if skills:
        feedback.append({
            "section": "Skills",
            "status": "Detected",
            "message": "Skills section is detected."
        })
    else:
        feedback.append({
            "section": "Skills",
            "status": "Missing",
            "message": "Skills section is not clearly detected. Consider adding relevant technical and soft skills."
        })

    if educations:
        feedback.append({
            "section": "Education",
            "status": "Detected",
            "message": "Education section is detected."
        })
    else:
        feedback.append({
            "section": "Education",
            "status": "Missing",
            "message": "Education section is not clearly detected. Consider adding your school, degree, and study period."
        })

    if experiences:
        feedback.append({
            "section": "Experience",
            "status": "Detected",
            "message": "Experience section is detected."
        })
    else:
        feedback.append({
            "section": "Experience",
            "status": "Missing",
            "message": "Experience section is not clearly detected. Consider adding work experience, internship, or relevant responsibilities."
        })

    if projects:
        feedback.append({
            "section": "Projects",
            "status": "Detected",
            "message": "Projects section is detected."
        })
    else:
        feedback.append({
            "section": "Projects",
            "status": "Missing",
            "message": "Projects section is not clearly detected. Consider adding academic or personal projects."
        })

    if certifications:
        feedback.append({
            "section": "Certifications",
            "status": "Detected",
            "message": "Certifications section is detected."
        })
    else:
        feedback.append({
            "section": "Certifications",
            "status": "Missing",
            "message": "Certifications section is not clearly detected. Consider adding relevant certificates if available."
        })

    return feedback

@app.route("/download_resume_docx", methods=["POST"])
def download_resume_docx():
    full_name = request.form.get("full_name", "Resume")
    email = request.form.get("email", "")
    phone = request.form.get("phone", "")
    location = request.form.get("location", "")
    target_role = request.form.get("target_role", "")

    summary = request.form.get("summary", "")
    skills = request.form.get("skills", "")
    experience = request.form.get("experience", "")
    education = request.form.get("education", "")
    projects = request.form.get("projects", "")
    certifications = request.form.get("certifications", "")

    document = docx.Document()

    document.add_heading(full_name, level=0)

    contact_line = " | ".join(
        item for item in [email, phone, location] if item
    )
    if contact_line:
        document.add_paragraph(contact_line)

    if target_role:
        document.add_paragraph(f"Target Role: {target_role}")

    if summary:
        document.add_heading("Professional Summary", level=1)
        document.add_paragraph(summary)

    if skills:
        document.add_heading("Skills", level=1)
        document.add_paragraph(skills)

    if experience:
        document.add_heading("Work Experience", level=1)
        document.add_paragraph(experience)

    if education:
        document.add_heading("Education", level=1)
        document.add_paragraph(education)

    if projects:
        document.add_heading("Projects", level=1)
        document.add_paragraph(projects)

    if certifications:
        document.add_heading("Certifications", level=1)
        document.add_paragraph(certifications)

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    filename = f"{full_name.replace(' ', '_')}_Resume.docx"

    return send_file(
        file_stream,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def safe_json_loads(value):
    try:
        return json.loads(value) if value else []
    except json.JSONDecodeError:
        return []

def build_resume_text_from_builder(resume_data):
    resume_parts = []

    resume_parts.append(resume_data.get("summary", ""))
    resume_parts.append(resume_data.get("ai_summary", ""))

    for skill in resume_data.get("skills", []):
        resume_parts.append(skill)

    for exp in resume_data.get("experiences", []):
        resume_parts.append(exp.get("company", ""))
        resume_parts.append(exp.get("role", ""))
        resume_parts.append(exp.get("period", ""))
        resume_parts.append(exp.get("description", ""))
        resume_parts.append(exp.get("ai_bullets", ""))

    for edu in resume_data.get("educations", []):
        resume_parts.append(edu.get("school", ""))
        resume_parts.append(edu.get("degree", ""))
        resume_parts.append(edu.get("period", ""))

    for project in resume_data.get("projects", []):
        resume_parts.append(project.get("title", ""))
        resume_parts.append(project.get("tools", ""))
        resume_parts.append(project.get("description", ""))

    for cert in resume_data.get("certifications", []):
        resume_parts.append(cert.get("name", ""))
        resume_parts.append(cert.get("provider", ""))
        resume_parts.append(cert.get("year", ""))

    return " ".join(resume_parts).lower()

def generate_ai_resume_summary(resume_data):
    prompt = f"""
    You are an AI resume writing assistant.

    Candidate Information:
    Name: {resume_data.get("full_name")}
    Target Role: {resume_data.get("target_role")}
    Current Summary: {resume_data.get("summary")}
    Skills: {", ".join(resume_data.get("skills", []))}
    Experience: {resume_data.get("experience_text")}
    Education: {resume_data.get("education_text")}
    Projects: {resume_data.get("project_text")}
    Certifications: {resume_data.get("certification_text")}

    Task:
    Rewrite the professional summary into an ATS-friendly resume summary.

    Requirements:
    1. Keep it around 2-3 sentences.
    2. Use professional language.
    3. Highlight relevant skills, education, projects, or experience.
    4. Do not invent fake experience, fake skills, or fake certifications.
    5. Make it suitable for the target role.
    """

    try:
        response = model.generate_content(prompt)
        return response.text.strip()

    except exceptions.ResourceExhausted:
        return resume_data.get("summary", "")

    except Exception as e:
        print(f"AI resume summary error: {e}")
        return resume_data.get("summary", "")
    
def generate_ai_bullet_points(experience):
    prompt = f"""
    You are an AI resume writing assistant.

    Convert the following work experience into 3 professional resume bullet points.

    Company: {experience.get("company")}
    Role: {experience.get("role")}
    Period: {experience.get("period")}
    What the candidate did:
    {experience.get("description")}

    Requirements:
    1. Use strong action verbs.
    2. Make the bullets suitable for an ATS-friendly resume.
    3. Do not invent fake achievements or numbers.
    4. Keep each bullet concise.
    """

    try:
        response = model.generate_content(prompt)
        return response.text.strip()

    except exceptions.ResourceExhausted:
        return experience.get("description", "")

    except Exception as e:
        print(f"AI bullet generation error: {e}")
        return experience.get("description", "")

def generate_builder_ai_feedback(score, matched, missing, ats_feedback, target_role):
    missing_sections = [
        item["section"] for item in ats_feedback if item["status"] == "Missing"
    ]

    prompt = f"""
    You are an AI resume advisor.

    Target Role: {target_role}
    Resume Score: {score}%

    Matched Skills:
    {", ".join(matched) if matched else "No strong matched skills detected."}

    Missing Skills:
    {", ".join(missing) if missing else "No major missing skills detected."}

    Missing ATS Sections:
    {", ".join(missing_sections) if missing_sections else "No major missing sections detected."}

    Task:
    Give a clear explanation of the resume score and suggest what the job seeker should improve.

    Requirements:
    1. Explain what the score means.
    2. Mention strengths.
    3. Mention missing skills or skill gaps.
    4. Mention ATS structure issues if any.
    5. Give practical improvement suggestions.
    6. Use a helpful and honest tone.
    7. Write in 2 short paragraphs.
    """

    try:
        response = model.generate_content(prompt)
        return response.text.strip()

    except exceptions.ResourceExhausted:
        return "AI feedback is temporarily unavailable due to usage limits. Please review the matched skills, missing skills, and ATS structure feedback."

    except Exception as e:
        print(f"AI builder feedback error: {e}")
        return "AI feedback is currently unavailable. Please review the score, matched skills, missing skills, and ATS structure feedback."

@app.route("/generate_resume", methods=["POST"])
def generate_resume():
    skills = safe_json_loads(request.form.get("skills_json"))
    experiences = safe_json_loads(request.form.get("experiences_json"))
    educations = safe_json_loads(request.form.get("educations_json"))
    projects = safe_json_loads(request.form.get("projects_json"))
    certifications = safe_json_loads(request.form.get("certifications_json"))

    resume_data = {
        "full_name": request.form.get("full_name", "").strip(),
        "email": request.form.get("email", "").strip(),
        "phone": request.form.get("phone", "").strip(),
        "location": request.form.get("location", "").strip(),
        "target_role": request.form.get("target_role", "").strip(),
        "target_job_description": request.form.get("target_job_description", "").strip(),
        "summary": request.form.get("summary", "").strip(),
        "skills": skills,
        "experiences": experiences,
        "educations": educations,
        "projects": projects,
        "certifications": certifications
    }

    enhanced_experiences = []

    for exp in experiences:
        ai_bullets = generate_ai_bullet_points(exp)
        exp["ai_bullets"] = ai_bullets
        enhanced_experiences.append(exp)

    resume_data["experiences"] = enhanced_experiences

    resume_data["experience_text"] = "\n\n".join(
        [
            f"{exp.get('role', '')} at {exp.get('company', '')} ({exp.get('period', '')})\n{exp.get('ai_bullets', '')}"
            for exp in enhanced_experiences
        ]
    )

    resume_data["education_text"] = "\n".join(
        [
            f"{edu.get('degree', '')}, {edu.get('school', '')} ({edu.get('period', '')})"
            for edu in educations
        ]
    )

    resume_data["project_text"] = "\n\n".join(
        [
            f"{project.get('title', '')}\nTools: {project.get('tools', '')}\n{project.get('description', '')}"
            for project in projects
        ]
    )

    resume_data["certification_text"] = "\n".join(
        [
            f"{cert.get('name', '')}, {cert.get('provider', '')} ({cert.get('year', '')})"
            for cert in certifications
        ]
    )

    ai_summary = generate_ai_resume_summary(resume_data)
    resume_data["ai_summary"] = ai_summary

    resume_text = build_resume_text_from_builder(resume_data)

    target_job_description = resume_data.get("target_job_description", "")

    if target_job_description:
        required_skills = extract_required_skills(target_job_description)

        score, matched, missing = calculate_combined_score(
            resume_text,
            target_job_description,
            required_skills
        )
    else:
        required_skills = []
        score = 0
        matched = []
        missing = []

    ats_feedback = check_builder_ats_structure(resume_data)

    ai_feedback = generate_builder_ai_feedback(
        score=score,
        matched=matched,
        missing=missing,
        ats_feedback=ats_feedback,
        target_role=resume_data.get("target_role", "")
    )

    return render_template(
        "resume_preview.html",
        resume=resume_data,
        score=score,
        matched=matched,
        missing=missing,
        required_skills=sorted(required_skills),
        ats_feedback=ats_feedback,
        ai_feedback=ai_feedback
    )

@app.route("/service-worker.js")
def service_worker():
    return send_from_directory("static", "service-worker.js")

@app.route("/uploads/<path:filename>")
def uploaded_file(filename):
    return send_from_directory(UPLOAD_DIR, filename)

if __name__ == "__main__":
    app.run(debug=True)
